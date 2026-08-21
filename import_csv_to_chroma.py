#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""离线入库脚本(独立运行, 不启动任何 web 服务).

数据流: docs/file.csv → csv 解析 → LangChain Document → 文本切片
       → AliDashScopeEmbeddings(text-embedding-v3) 向量化 → Chroma 本地持久化

用法:
    python import_csv_to_chroma.py            # 增量 upsert(按稳定 ID 幂等, 可重复执行)
    python import_csv_to_chroma.py --rebuild  # 清空集合后全量重建索引
"""
import argparse
import csv
import hashlib
import io
import json
import logging
import os
import sys
import time
from pathlib import Path
from typing import List, Tuple

from dotenv import load_dotenv

load_dotenv()

from chromadb.config import Settings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from embedding import AliDashScopeEmbeddings

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger("ingest")

CSV_PATH = os.getenv("CSV_PATH", "docs/file.csv")
CHROMA_DIR = os.getenv("CHROMA_DIR", "chroma_db")
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "rag_pro_kb")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-v3")
EMBEDDING_BASE_URL = os.getenv("EMBEDDING_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1")
EMBEDDING_API_KEY = os.getenv("DASHSCOPE_API_KEY", "")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "500"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "50"))


def parse_answer_field(raw: str) -> str:
    """答案字段形如 JSON 数组字符串 ["xxx"], 取首个非空元素; 解析失败时返回原文。"""
    if not raw or not raw.strip():
        return ""
    raw = raw.strip()
    try:
        data = json.loads(raw)
        if isinstance(data, list):
            for item in data:
                if isinstance(item, str) and item.strip():
                    return item.strip()
            return ""
        if isinstance(data, str) and data.strip():
            return data.strip()
    except (json.JSONDecodeError, ValueError):
        pass
    return raw.strip().strip('"').strip()


def load_rows(csv_path: str) -> List[dict]:
    """读取 csv 文件, 每一行封装为一条知识素材(question/answer), 过滤空值。"""
    with open(csv_path, "r", encoding="utf-8-sig", newline="") as f:
        return load_rows_from_text(f.read())


def xlsx_to_csv_text(data: bytes) -> str:
    """读取 xlsx 第一个 sheet, 转为 csv 文本以复用 csv 解析流程。"""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        ws = wb.worksheets[0]
        buf = io.StringIO(newline="")
        writer = csv.writer(buf)
        for row in ws.iter_rows(values_only=True):
            if row:  # 跳过整行为空的数据
                writer.writerow(["" if v is None else str(v).strip() for v in row])
        return buf.getvalue()
    finally:
        wb.close()


def docx_to_text(data: bytes) -> str:
    """读取 docx 文档纯文本(段落与表格文本按换行拼接)。"""
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # 表格内文本同样纳入, 避免丢失结构化内容
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def pdf_to_text(data: bytes) -> str:
    """读取 pdf 文档纯文本(按页提取并拼接)。"""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())
    return "\n".join(parts)


def load_rows_from_text(text: str) -> List[dict]:
    """解析 csv 文本内容, 每一行封装为一条知识素材(question/answer), 过滤空值。"""
    rows: List[dict] = []
    reader = csv.reader(io.StringIO(text, newline=""))
    header = next(reader, None)
    if not header:
        return rows
    idx = {name.strip(): i for i, name in enumerate(header)}
    id_i = idx.get("id", 0)
    q_i = idx.get("question", 1)
    h_i = idx.get("human_answers", 2)
    c_i = idx.get("chatgpt_answers", 3)

    for line_no, row in enumerate(reader, start=2):
        if not row or all(not cell.strip() for cell in row):
            continue  # 空行过滤
        if len(row) < 4:
            row = row + [""] * (4 - len(row))  # 兼容缺字段的行
        question = row[q_i].strip()
        # 优先人工答案, 缺失时回退 chatgpt 答案
        answer = parse_answer_field(row[h_i]) or parse_answer_field(row[c_i])
        if not question and not answer:
            log.warning("第 %d 行问题与答案均为空, 已过滤", line_no)
            continue
        rows.append(
            {
                "row_id": row[id_i].strip() or str(line_no),
                "question": question,
                "answer": answer,
                "line_no": line_no,
            }
        )
    return rows


def _make_splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "；", "！", "？", "，", " ", ""],
    )


def build_documents(rows: List[dict], source: str = None) -> Tuple[List[Document], List[str]]:
    """将 csv 解析出的知识行封装为 LangChain Document 并切片, 返回 (docs, 稳定ID列表)。"""
    source = source or CSV_PATH
    splitter = _make_splitter()
    docs: List[Document] = []
    ids: List[str] = []
    for r in rows:
        if r["answer"]:
            # 每个切片前置原始问题, 提升向量检索的召回质量
            contents = [f"问题：{r['question']}\n内容：{c}" for c in splitter.split_text(r["answer"])]
        else:
            contents = [r["question"]]  # 无答案的行(如个人简介类)直接以问题文本入库
        for i, content in enumerate(contents):
            docs.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": source,
                        "row_id": r["row_id"],
                        "question": r["question"],
                        "chunk_index": i,
                    },
                )
            )
            ids.append(hashlib.md5(f"{source}|{r['row_id']}|{r['question']}|{i}".encode("utf-8")).hexdigest())
    return docs, ids


def build_text_documents(filename: str, text: str) -> Tuple[List[Document], List[str]]:
    """将纯文本(txt/md)切片封装为 LangChain Document, 返回 (docs, 稳定ID列表)。"""
    splitter = _make_splitter()
    docs: List[Document] = []
    ids: List[str] = []
    for i, chunk in enumerate(splitter.split_text(text.strip())):
        docs.append(
            Document(
                page_content=chunk,
                metadata={"source": filename, "row_id": filename, "question": filename, "chunk_index": i},
            )
        )
        ids.append(hashlib.md5(f"{filename}|upload|{i}".encode("utf-8")).hexdigest())
    return docs, ids


def main() -> None:
    parser = argparse.ArgumentParser(description="CSV 知识库离线入库到 Chroma")
    parser.add_argument("--rebuild", action="store_true", help="清空集合并全量重建索引")
    parser.add_argument("--csv", default=CSV_PATH, help=f"csv 路径, 默认 {CSV_PATH}")
    args = parser.parse_args()

    if not EMBEDDING_API_KEY or "请替换" in EMBEDDING_API_KEY:
        log.error("缺少 DASHSCOPE_API_KEY, 请在 .env 中配置后重试")
        sys.exit(1)
    if not Path(args.csv).exists():
        log.error("csv 文件不存在: %s", args.csv)
        sys.exit(1)

    rows = load_rows(args.csv)
    log.info("csv 解析完成: %s, 有效知识条目 %d 条", args.csv, len(rows))
    docs, ids = build_documents(rows, source=args.csv)
    log.info("切片完成: 共 %d 个切片 (chunk_size=%d, overlap=%d)", len(docs), CHUNK_SIZE, CHUNK_OVERLAP)

    embeddings = AliDashScopeEmbeddings(
        api_key=EMBEDDING_API_KEY, base_url=EMBEDDING_BASE_URL, model=EMBEDDING_MODEL
    )

    def _make_vectorstore() -> Chroma:
        return Chroma(
            collection_name=COLLECTION_NAME,
            embedding_function=embeddings,
            persist_directory=CHROMA_DIR,
            collection_metadata={"hnsw:space": "cosine"},  # 余弦相似度
            client_settings=Settings(anonymized_telemetry=False, is_persistent=True),
        )

    vectorstore = _make_vectorstore()
    if args.rebuild:
        try:
            vectorstore.delete_collection()
            log.info("已删除旧集合 %s, 开始全量重建", COLLECTION_NAME)
        except Exception:  # noqa: BLE001 集合不存在时忽略
            pass
        vectorstore = _make_vectorstore()

    start = time.time()
    batch = 32
    for i in range(0, len(docs), batch):
        part = docs[i : i + batch]
        vectorstore.add_documents(part, ids=ids[i : i + batch])
        log.info("已入库 %d/%d", min(i + batch, len(docs)), len(docs))
    log.info(
        "入库完成: 集合=%s, 共 %d 条, 耗时 %.1fs, 持久化目录=%s",
        COLLECTION_NAME, len(docs), time.time() - start, CHROMA_DIR,
    )


if __name__ == "__main__":
    main()
